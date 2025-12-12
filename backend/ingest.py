"""
Document ingestion - extract text from PDFs, DOCX, TXT, Markdown with improved error handling.
"""

import os
from pathlib import Path
from typing import List, Tuple, Optional
import re

from logger_config import logger


class DocumentIngestor:
    """Handle document extraction and chunking with proper error handling."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document ingestor.
        
        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
        """
        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive")
        if chunk_overlap < 0:
            raise ValueError("chunk_overlap must be non-negative")
        if chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be less than chunk_size")
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"DocumentIngestor initialized: chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    def load_and_process_documents(self, file_paths: List[str]) -> Tuple[List[str], str]:
        """
        Load and chunk documents from file paths.
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            Tuple of (chunks, combined_document_name)
        """
        if not file_paths:
            logger.warning("No file paths provided")
            return [], "unknown"
        
        all_text = ""
        doc_names = []
        
        for file_path in file_paths:
            if not os.path.exists(file_path):
                logger.warning(f"File not found: {file_path}")
                continue
            
            try:
                text = self._extract_text(file_path)
                if text and text.strip():
                    all_text += "\n\n" + text
                    doc_names.append(Path(file_path).stem)
                    logger.debug(f"Extracted text from {file_path}: {len(text)} characters")
                else:
                    logger.warning(f"No text extracted from {file_path}")
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}", exc_info=True)
                continue
        
        if not all_text.strip():
            logger.warning("No text extracted from any files")
            return [], "unknown"
        
        chunks = self._chunk_text(all_text)
        doc_name = " + ".join(doc_names) if doc_names else "unknown"
        
        logger.info(f"Processed {len(file_paths)} document(s) into {len(chunks)} chunks")
        return chunks, doc_name
    
    def _extract_text(self, file_path: str) -> str:
        """
        Extract text based on file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text
        """
        file_path_lower = str(file_path).lower()
        
        if file_path_lower.endswith('.pdf'):
            return self._extract_pdf(file_path)
        elif file_path_lower.endswith('.docx'):
            return self._extract_docx(file_path)
        elif file_path_lower.endswith(('.txt', '.md')):
            return self._extract_text_file(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_path}")
            return ""
    
    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            from PyPDF2 import PdfReader
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num} from {file_path}: {e}")
                        continue
                
                logger.debug(f"Extracted {num_pages} pages from PDF: {file_path}")
            
            return text.strip()
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}", exc_info=True)
            return ""
    
    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            text = ""
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            logger.debug(f"Extracted text from DOCX: {file_path}")
            return text.strip()
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}", exc_info=True)
            return ""
    
    def _extract_text_file(self, file_path: str) -> str:
        """
        Extract text from TXT or Markdown file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File contents
        """
        try:
            # Try UTF-8 first, fallback to other encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        logger.debug(f"Extracted text from {file_path} using {encoding}")
                        return content
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode {file_path} with any encoding")
            return ""
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}", exc_info=True)
            return ""
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks with sentence boundary awareness.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                last_period = text.rfind('.', start, end)
                last_exclamation = text.rfind('!', start, end)
                last_question = text.rfind('?', start, end)
                
                # Use the latest sentence boundary
                sentence_end = max(last_period, last_exclamation, last_question)
                
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Fallback to word boundary
                    last_space = text.rfind(' ', start, end)
                    if last_space > start:
                        end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position accounting for overlap
            start = max(start + 1, end - self.chunk_overlap)
            
            # Safety check to prevent infinite loop
            if start >= len(text):
                break
        
        logger.debug(f"Chunked text into {len(chunks)} chunks")
        return chunks
    
    def process_uploaded_file(self, file_content: bytes, filename: str) -> Tuple[List[str], str]:
        """
        Process uploaded file from bytes.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (chunks, document_name)
        """
        if not file_content:
            logger.warning("Empty file content provided")
            return [], "unknown"
        
        temp_path = Path("data/documents") / filename
        temp_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            # Write temporary file
            with open(temp_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"Processing uploaded file: {filename} ({len(file_content)} bytes)")
            chunks, doc_name = self.load_and_process_documents([str(temp_path)])
            return chunks, doc_name
            
        except Exception as e:
            logger.error(f"Error processing uploaded file {filename}: {e}", exc_info=True)
            return [], filename or "unknown"
        finally:
            # Clean up temporary file
            try:
                if temp_path.exists():
                    temp_path.unlink()
                    logger.debug(f"Cleaned up temporary file: {temp_path}")
            except Exception as e:
                logger.warning(f"Error cleaning up temporary file {temp_path}: {e}")
